#!/usr/bin/env python3
"""Genera documentos Word de FORMULARIO WCAG 2.2 para el grupo de desarrollo."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "scripts" / "formulario-data.json"
OUT_DIR = ROOT / "docs" / "formularios"
EVIDENCIA = ROOT / "docs" / "evidencia"


def load_data() -> dict:
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_field(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)


def build_criterion_block(doc: Document, c: dict, sistema: str) -> None:
    add_heading(doc, f"{c['id']} — {c['nombre']}", level=2)
    add_field(doc, "Sistema", sistema)
    add_field(doc, "Nivel WCAG", c["nivel"])
    add_field(doc, "Estado", c["estado"])
    add_field(doc, "Opción del menú de accesibilidad", c["menu"])

    p = doc.add_paragraph()
    p.add_run("Descripción de conformidad\n").bold = True
    doc.add_paragraph(c["formulario"])

    add_field(doc, "Evidencia técnica", c["evidencia"])
    add_field(doc, "Captura recomendada", c["captura"])
    doc.add_paragraph()


def add_cover(doc: Document, title: str, subtitle: str, meta: dict) -> None:
    add_heading(doc, title, level=0)
    doc.add_paragraph(subtitle)
    doc.add_paragraph()
    add_field(doc, "Sistema", meta["sistema"])
    add_field(doc, "Fecha", meta["fecha"])
    doc.add_paragraph()


def embed_evidencia_if_exists(doc: Document, names: list[str]) -> None:
    add_heading(doc, "Anexos — capturas de evidencia", level=1)
    for name in names:
        path = EVIDENCIA / name
        if not path.exists():
            continue
        doc.add_paragraph(name)
        doc.add_picture(str(path), width=Inches(5.5))
        doc.add_paragraph()


def generate_complete(data: dict) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(11)

    add_cover(
        doc,
        "FORMULARIO WCAG 2.2 A/AA — 55 criterios",
        "Sistema Inteligente de Orientación Vocacional (1.ª prueba)",
        data,
    )

    add_heading(doc, "Resumen", level=1)
    estados: dict[str, int] = {}
    for c in data["criterios"]:
        estados[c["estado"]] = estados.get(c["estado"], 0) + 1
    for k, v in sorted(estados.items()):
        add_field(doc, k, str(v))
    add_field(doc, "Total", str(len(data["criterios"])))
    doc.add_page_break()

    categories = [
        ("1. Alternativas textuales y multimedia", lambda x: x.startswith("1.1") or x.startswith("1.2")),
        ("2. Adaptable — estructura e idioma", lambda x: x.startswith("1.3") or x.startswith("3.1")),
        ("3. Distinguible — color y contraste", lambda x: x.startswith("1.4")),
        ("4. Operable — teclado y navegación", lambda x: x.startswith("2.1") or x.startswith("2.4")),
        ("5. Tiempo y movimiento", lambda x: x.startswith("2.2") or x.startswith("2.3")),
        ("6. Puntero y táctil", lambda x: x.startswith("2.5")),
        ("7. Comprensible — predecible", lambda x: x.startswith("3.2")),
        ("8. Formularios y errores", lambda x: x.startswith("3.3")),
        ("9. Robusto", lambda x: x.startswith("4.1")),
    ]

    for title, pred in categories:
        items = [c for c in data["criterios"] if pred(c["id"])]
        if not items:
            continue
        add_heading(doc, title, level=1)
        for c in items:
            build_criterion_block(doc, c, data["sistema"])

    embed_evidencia_if_exists(
        doc,
        [
            "landing.png",
            "ayuda.png",
            "login.png",
            "menu-accesibilidad.png",
            "menu-cognitivo.png",
        ],
    )

    out = OUT_DIR / "FORMULARIO_Completo_55criterios.docx"
    doc.save(out)
    return out


def generate_member(data: dict, key: str) -> Path:
    member = data["integrantes"][key]
    doc = Document()
    add_cover(
        doc,
        f"FORMULARIO — {member['nombre']}",
        "Checklist grupo desarrollo · WCAG 2.2 A/AA",
        data,
    )

    if member["criterios"] == "resto":
        assigned = [
            c
            for c in data["criterios"]
            if c["id"] not in data["integrantes"]["jeremy"]["criterios"]
            and c["id"] not in data["integrantes"]["justin"]["criterios"]
        ]
    else:
        assigned = [c for c in data["criterios"] if c["id"] in member["criterios"]]

    add_field(doc, "Criterios asignados", str(len(assigned)))
    doc.add_paragraph()

    for c in assigned:
        build_criterion_block(doc, c, data["sistema"])

    if key == "justin":
        embed_evidencia_if_exists(doc, ["landing.png"])
    elif key == "jeremy":
        embed_evidencia_if_exists(doc, ["landing.png"])
    elif key == "cesar":
        embed_evidencia_if_exists(
            doc,
            [
                "ayuda.png",
                "login.png",
                "menu-accesibilidad.png",
                "menu-cognitivo.png",
            ],
        )

    slug = key.replace("_", "-")
    out = OUT_DIR / f"FORMULARIO_{slug}.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    files = [
        generate_complete(data),
        generate_member(data, "jeremy"),
        generate_member(data, "justin"),
        generate_member(data, "cesar"),
    ]
    print("Generados:")
    for f in files:
        print(f"  {f.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
